import asyncio
import os
import shutil
import tempfile
import unittest

import pandas
from docx import Document

from dbscripts import BotDatabase
from expdata import ExportData


class TestExportData(unittest.TestCase):
    def setUp(self):
        temp_dir = tempfile.mkdtemp()
        self.db_path = os.path.join(temp_dir, 'test.db')

        self.database = BotDatabase(self.db_path)
        self.database.create_tables()

        # Запускаем метод save_user_data
        asyncio.run(self.database.save_user_data(
            "test_id",
            "test_user_name",
            "test_request_id",
            "test_problem_description",
            "test_contact_info",
            "test_contact_time"
        ))

    def tearDown(self):
        shutil.rmtree(os.path.dirname(self.db_path))

    def test_export_to_word(self):
        docx_path = os.path.join(os.path.dirname(self.db_path), 'test.docx')

        ExportData.export_to_word(self.db_path, 'users', output_file=docx_path)

        self.assertTrue(os.path.exists(docx_path))

        doc = Document(docx_path)
        if doc.tables:
            table = doc.tables[0]  # Получаем первую таблицу
            rows = table.rows[1]  # Получаем вторую строку
            data = [cell.text for cell in rows.cells]  # Получаем текст из каждой ячейки

            self.assertEqual(['test_id', 'test_user_name', 'test_contact_info'], data)

    def test_export_to_excel(self):
        xlsx_path = os.path.join(os.path.dirname(self.db_path), 'test.xlsx')
        ExportData.export_to_excel(self.db_path, 'users', output_file=xlsx_path)

        self.assertTrue(os.path.exists(xlsx_path))

        table = pandas.read_excel(xlsx_path)
        data = table.iloc[0].to_list()

        self.assertEqual(['test_id', 'test_user_name', 'test_contact_info'], data)

    def test_export_to_csv(self):
        csv_path = os.path.join(os.path.dirname(self.db_path), 'test.csv')
        ExportData.export_to_csv(self.db_path, 'users', output_file=csv_path)

        self.assertTrue(os.path.exists(csv_path))

        table = pandas.read_csv(csv_path)
        data = table.iloc[0].to_list()

        self.assertEqual(['test_id', 'test_user_name', 'test_contact_info'], data)

    def test_export_to_html(self):
        html_path = os.path.join(os.path.dirname(self.db_path), 'test.html')
        ExportData.export_to_html(self.db_path, 'users', output_file=html_path)

        self.assertTrue(os.path.exists(html_path))

        table = pandas.read_html(html_path)
        data = table[0]

        self.assertEqual(['test_id', 'test_user_name', 'test_contact_info'], data.iloc[0].tolist())



